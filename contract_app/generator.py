"""Word 템플릿의 {{변수}}를 사용자 답안으로 치환해 .docx 생성.

이슈: python-docx로 만든 templates/*.docx는 한 paragraph 내부에서 단일 run에 텍스트가
모여 있어 단순 r.text 치환으로 충분하지만, 외부 환경에서 만든 docx는 run 단위가
잘게 쪼개져 {{변수}} 토큰이 여러 run에 분산될 수 있다. 그래서 우리는 paragraph 단위로
전체 텍스트를 모은 뒤 첫 run에 결과를 쓰고 나머지 run은 비워서 치환한다.
"""
from __future__ import annotations

import io
import re
from copy import deepcopy
from pathlib import Path
from typing import Mapping

from docx import Document
from docx.text.paragraph import Paragraph


VAR_RE = re.compile(r"\{\{\s*([^\}\s]+)\s*\}\}")


def _replace_in_paragraph(p: Paragraph, mapping: Mapping[str, str]) -> int:
    """paragraph 안의 {{var}}들을 치환. 치환된 변수 수 반환.
    run 단위 분산을 처리하기 위해 모든 run을 합친 텍스트에서 정규식 치환 후
    첫 run에 결과를 몰아넣고 나머지 run의 text를 비운다.
    """
    if not p.runs:
        return 0
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return 0
    count = [0]
    def sub(m: re.Match[str]) -> str:
        var = m.group(1)
        if var in mapping:
            count[0] += 1
            return str(mapping[var])
        return m.group(0)
    new_text = VAR_RE.sub(sub, full)
    if new_text == full:
        return 0
    # 서식 보존: 첫 run에 새 텍스트, 나머지 run 텍스트는 비우기
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""
    return count[0]


def _iter_all_paragraphs(doc) -> list[Paragraph]:
    """문서 내 모든 paragraph(테이블 셀 내부 포함, 헤더/푸터 포함)."""
    ps: list[Paragraph] = []
    ps.extend(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                ps.extend(cell.paragraphs)
                for inner in cell.tables:
                    for r2 in inner.rows:
                        for c2 in r2.cells:
                            ps.extend(c2.paragraphs)
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header,
                     section.footer, section.first_page_footer, section.even_page_footer):
            if part is None:
                continue
            ps.extend(part.paragraphs)
            for t in part.tables:
                for row in t.rows:
                    for cell in row.cells:
                        ps.extend(cell.paragraphs)
    return ps


def render(template_path: Path, mapping: Mapping[str, str]) -> bytes:
    """템플릿에 mapping을 적용한 docx 바이트 반환."""
    doc = Document(str(template_path))
    total = 0
    for p in _iter_all_paragraphs(doc):
        total += _replace_in_paragraph(p, mapping)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def find_unfilled(template_path: Path, mapping: Mapping[str, str]) -> list[str]:
    """렌더 후에도 치환되지 않은 {{변수}}가 있는지 확인."""
    doc = Document(str(template_path))
    leftover: set[str] = set()
    for p in _iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        for m in VAR_RE.finditer(full):
            if m.group(1) not in mapping:
                leftover.add(m.group(1))
    return sorted(leftover)


def list_template_vars(template_path: Path) -> list[str]:
    """템플릿에 존재하는 {{변수}} 전체 목록."""
    doc = Document(str(template_path))
    found: set[str] = set()
    for p in _iter_all_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        for m in VAR_RE.finditer(full):
            found.add(m.group(1))
    return sorted(found)
